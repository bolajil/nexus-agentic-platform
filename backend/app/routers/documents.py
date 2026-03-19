"""
NEXUS Platform — Document Upload Router
========================================
Accepts internal project documents (PDF, DOCX, TXT, MD) and ingests them
into the ChromaDB vector knowledge base with intelligent chunking,
terminology normalization, and project-scoped metadata.

Endpoints:
  POST /documents/upload          — Upload file(s) for a project
  GET  /documents                 — List all ingested documents
  GET  /documents/{doc_id}        — Get document metadata
  DELETE /documents/{doc_id}      — Remove document from KB

Why chunking matters for RAG:
  Embedding a full 50-page PDF as one vector loses granularity.
  A query about "fin efficiency" should retrieve the 3-paragraph
  section that defines η_f = tanh(mL)/(mL), not the entire document.

  We use RecursiveCharacterTextSplitter with:
    - chunk_size = 800 tokens (fits in LLM context with headroom)
    - chunk_overlap = 120 tokens (preserves cross-boundary context)
    - separators = ["\n\n", "\n", ". ", " "] (respects natural breaks)

Why terminology normalization at index time:
  If a document says "θ_ja = 0.6°C/W" and the user queries
  "junction-to-ambient resistance", pure cosine similarity may miss it
  because the embedding model tokenizes "θ_ja" differently from the
  spelled-out form. We normalize BOTH the document chunks AND queries
  before embedding, ensuring they land in the same semantic space.
"""
from __future__ import annotations

import hashlib
import logging
import mimetypes
import uuid
from datetime import datetime
from pathlib import Path
from typing import Optional

from fastapi import APIRouter, Depends, File, Form, HTTPException, Response, UploadFile

from app.core.config import Settings, get_settings
from app.memory.vector_store import VectorStoreManager
from app.tools.terminology import normalize_document, register_project_glossary, extract_technical_terms

logger = logging.getLogger(__name__)
router = APIRouter(prefix="/documents", tags=["documents"])

# Supported MIME types → extractor function name
SUPPORTED_TYPES = {
    "application/pdf":   "_extract_pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "_extract_docx",
    "text/plain":        "_extract_text",
    "text/markdown":     "_extract_text",
    "text/x-markdown":   "_extract_text",
}
MAX_FILE_SIZE_MB = 50
CHUNK_SIZE = 800        # characters (approx 200 tokens)
CHUNK_OVERLAP = 120     # characters of overlap between chunks

_vector_store: Optional[VectorStoreManager] = None


def get_vector_store(settings: Settings = Depends(get_settings)) -> VectorStoreManager:
    global _vector_store
    if _vector_store is None:
        _vector_store = VectorStoreManager(
            openai_api_key=settings.openai_api_key,
            host=settings.chroma_host,
            port=settings.chroma_port,
        )
        _vector_store.initialize()
    return _vector_store


# ── Text Extractors ───────────────────────────────────────────────────────────

def _extract_pdf(data: bytes) -> str:
    """Extract text from PDF using PyMuPDF (fitz) or pdfminer fallback."""
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(stream=data, filetype="pdf")
        pages = []
        for page in doc:
            text = page.get_text("text")
            if text.strip():
                pages.append(text)
        doc.close()
        return "\n\n".join(pages)
    except ImportError:
        pass

    try:
        from pdfminer.high_level import extract_text_to_fp
        from pdfminer.layout import LAParams
        import io
        output = io.StringIO()
        extract_text_to_fp(
            io.BytesIO(data), output,
            laparams=LAParams(), output_type="text", codec="utf-8"
        )
        return output.getvalue()
    except ImportError:
        raise HTTPException(
            status_code=422,
            detail="PDF extraction requires 'pymupdf' or 'pdfminer.six'. "
                   "Install with: pip install pymupdf"
        )


def _extract_docx(data: bytes) -> str:
    """Extract text from DOCX using python-docx."""
    try:
        import docx
        import io
        doc = docx.Document(io.BytesIO(data))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)
        return "\n\n".join(paragraphs)
    except ImportError:
        raise HTTPException(
            status_code=422,
            detail="DOCX extraction requires 'python-docx'. "
                   "Install with: pip install python-docx"
        )


def _extract_text(data: bytes) -> str:
    """Extract text from plain text / markdown files."""
    for encoding in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


# ── Chunking ─────────────────────────────────────────────────────────────────

def _chunk_text(text: str, chunk_size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> list[str]:
    """
    Recursive character text splitter.

    Tries to split on paragraph breaks (\\n\\n) first, then single newlines,
    then sentence boundaries ('. '), then spaces. This preserves natural
    semantic boundaries in engineering documents.

    Overlap ensures that context at chunk boundaries is not lost — critical
    when a formula spans the end of one paragraph and the start of the next.
    """
    separators = ["\n\n", "\n", ". ", " ", ""]

    def _split(text: str, seps: list[str]) -> list[str]:
        if not seps or len(text) <= chunk_size:
            return [text] if text.strip() else []

        sep = seps[0]
        parts = text.split(sep)
        chunks = []
        current = ""

        for part in parts:
            candidate = current + (sep if current else "") + part
            if len(candidate) <= chunk_size:
                current = candidate
            else:
                if current.strip():
                    chunks.append(current.strip())
                # If single part is too large, recurse with next separator
                if len(part) > chunk_size:
                    chunks.extend(_split(part, seps[1:]))
                    current = ""
                else:
                    current = part

        if current.strip():
            chunks.append(current.strip())

        return chunks

    raw_chunks = _split(text, separators)

    # Add overlap by prepending the tail of the previous chunk
    overlapped = []
    for i, chunk in enumerate(raw_chunks):
        if i > 0 and overlap > 0:
            prev_tail = raw_chunks[i - 1][-overlap:]
            chunk = prev_tail.strip() + " " + chunk
        overlapped.append(chunk)

    return overlapped


# ── Upload Endpoint ───────────────────────────────────────────────────────────

@router.post("/upload", status_code=201)
async def upload_document(
    file: UploadFile = File(...),
    domain: str = Form(..., description="Engineering domain: heat_transfer|propulsion|structural|electronics_cooling|general"),
    project_id: str = Form("default", description="Project identifier for scoped retrieval"),
    title: Optional[str] = Form(None, description="Document title (defaults to filename)"),
    source: Optional[str] = Form(None, description="Reference citation or source URL"),
    project_glossary: Optional[str] = Form(
        None,
        description=(
            "JSON string of project-specific terminology. "
            "Example: '{\"NEXUS-TPS\": \"NEXUS thermal protection system\", \"η_prop\": \"propulsive efficiency\"}'"
        ),
    ),
    store: VectorStoreManager = Depends(get_vector_store),
) -> dict:
    """
    Upload an internal project document (PDF, DOCX, TXT, MD) to the knowledge base.

    Processing pipeline:
      1. Validate file type and size
      2. Extract raw text (PDF/DOCX/TXT)
      3. Normalize terminology (Greek letters, abbreviations, project glossary)
      4. Split into overlapping chunks (800 chars, 120 overlap)
      5. Embed each chunk with text-embedding-3-small
      6. Store in ChromaDB with rich metadata (doc_id, chunk_index, domain, project_id)

    The project_id enables scoped retrieval — agents can filter by project
    so a query about "TPS" in Project A doesn't retrieve unrelated TPS
    docs from Project B.
    """
    # ── Validate file ──────────────────────────────────────────────────
    if not file.filename:
        raise HTTPException(status_code=400, detail="No filename provided")

    file_ext = Path(file.filename).suffix.lower()
    mime_type = file.content_type or mimetypes.guess_type(file.filename)[0] or ""

    # Normalise content-type for common mismatches
    if file_ext in (".md", ".markdown"):
        mime_type = "text/markdown"
    elif file_ext == ".txt":
        mime_type = "text/plain"

    if mime_type not in SUPPORTED_TYPES and file_ext not in (".pdf", ".docx", ".txt", ".md"):
        raise HTTPException(
            status_code=415,
            detail=f"Unsupported file type: {mime_type}. Supported: PDF, DOCX, TXT, MD",
        )

    data = await file.read()
    size_mb = len(data) / (1024 * 1024)
    if size_mb > MAX_FILE_SIZE_MB:
        raise HTTPException(status_code=413, detail=f"File too large: {size_mb:.1f}MB (max {MAX_FILE_SIZE_MB}MB)")

    # ── Extract text ───────────────────────────────────────────────────
    extractor_name = SUPPORTED_TYPES.get(mime_type, "_extract_text")
    if file_ext == ".pdf":
        extractor_name = "_extract_pdf"
    elif file_ext == ".docx":
        extractor_name = "_extract_docx"

    extractor = {
        "_extract_pdf":  _extract_pdf,
        "_extract_docx": _extract_docx,
        "_extract_text": _extract_text,
    }[extractor_name]

    raw_text = extractor(data)
    if not raw_text.strip():
        raise HTTPException(status_code=422, detail="Could not extract any text from the document")

    # ── Register project glossary ──────────────────────────────────────
    if project_glossary:
        try:
            import json
            glossary_dict = json.loads(project_glossary)
            register_project_glossary(project_id, glossary_dict)
        except (json.JSONDecodeError, ValueError) as e:
            raise HTTPException(status_code=400, detail=f"Invalid project_glossary JSON: {e}")

    # ── Normalize terminology ──────────────────────────────────────────
    normalized_text = normalize_document(raw_text, project_id=project_id)

    # ── Chunk ──────────────────────────────────────────────────────────
    chunks = _chunk_text(normalized_text)
    if not chunks:
        raise HTTPException(status_code=422, detail="Document produced no text chunks after processing")

    # ── Extract technical terms for metadata ──────────────────────────
    doc_terms = extract_technical_terms(normalized_text)

    # ── Build document records ─────────────────────────────────────────
    doc_id = hashlib.md5(f"{file.filename}{project_id}{datetime.utcnow().isoformat()}".encode()).hexdigest()[:16]
    doc_title = title or Path(file.filename).stem.replace("_", " ").replace("-", " ").title()

    documents = []
    for i, chunk in enumerate(chunks):
        chunk_id = f"{doc_id}_chunk_{i:04d}"
        documents.append({
            "id": chunk_id,
            "content": chunk,
            "metadata": {
                "doc_id":       doc_id,
                "chunk_index":  i,
                "total_chunks": len(chunks),
                "title":        doc_title,
                "filename":     file.filename,
                "domain":       domain,
                "project_id":   project_id,
                "source":       source or file.filename,
                "size_bytes":   len(data),
                "ingested_at":  datetime.utcnow().isoformat(),
                "technical_terms": ", ".join(doc_terms[:30]),  # store top 30 terms
            },
        })

    # ── Embed and store ────────────────────────────────────────────────
    n_added = store.add_documents(documents)

    logger.info(
        f"Uploaded '{file.filename}' → {n_added}/{len(chunks)} chunks ingested "
        f"[project={project_id}, domain={domain}]"
    )

    return {
        "doc_id":         doc_id,
        "title":          doc_title,
        "filename":       file.filename,
        "domain":         domain,
        "project_id":     project_id,
        "chunks_created": len(chunks),
        "chunks_ingested": n_added,
        "size_mb":        round(size_mb, 3),
        "characters":     len(raw_text),
        "technical_terms_found": doc_terms[:20],
        "status":         "ingested",
    }


@router.post("/upload/batch", status_code=201)
async def upload_batch(
    files: list[UploadFile] = File(...),
    domain: str = Form("general"),
    project_id: str = Form("default"),
    store: VectorStoreManager = Depends(get_vector_store),
) -> dict:
    """Upload multiple files at once. Returns per-file results."""
    results = []
    errors = []

    for file in files:
        try:
            # Reuse single upload logic
            result = await upload_document(
                file=file,
                domain=domain,
                project_id=project_id,
                store=store,
            )
            results.append(result)
        except HTTPException as e:
            errors.append({"filename": file.filename, "error": e.detail})
        except Exception as e:
            errors.append({"filename": file.filename, "error": str(e)})

    return {
        "total_files":  len(files),
        "succeeded":    len(results),
        "failed":       len(errors),
        "results":      results,
        "errors":       errors,
    }


@router.get("")
async def list_documents(
    project_id: Optional[str] = None,
    domain: Optional[str] = None,
    store: VectorStoreManager = Depends(get_vector_store),
) -> list[dict]:
    """
    List all ingested documents (deduplicated by doc_id).
    Optionally filter by project_id or domain.
    """
    try:
        # Query ChromaDB for chunk_index==0 records to get one entry per document
        where: dict = {"chunk_index": {"$eq": 0}}
        if project_id:
            where["project_id"] = {"$eq": project_id}
        if domain:
            where["domain"] = {"$eq": domain}

        results = store._collection.get(where=where, include=["metadatas"])
        docs = []
        for meta in (results.get("metadatas") or []):
            docs.append({
                "doc_id":       meta.get("doc_id"),
                "title":        meta.get("title"),
                "filename":     meta.get("filename"),
                "domain":       meta.get("domain"),
                "project_id":   meta.get("project_id"),
                "total_chunks": meta.get("total_chunks"),
                "ingested_at":  meta.get("ingested_at"),
                "size_bytes":   meta.get("size_bytes"),
            })
        return sorted(docs, key=lambda d: d.get("ingested_at", ""), reverse=True)
    except Exception as e:
        logger.error(f"List documents failed: {e}")
        return []


@router.delete("/{doc_id}")
async def delete_document(
    doc_id: str,
    store: VectorStoreManager = Depends(get_vector_store),
) -> Response:
    """Remove all chunks of a document from the knowledge base."""
    try:
        # Get all chunk IDs for this doc
        results = store._collection.get(
            where={"doc_id": {"$eq": doc_id}},
            include=["metadatas"],
        )
        ids_to_delete = results.get("ids", [])
        if not ids_to_delete:
            raise HTTPException(status_code=404, detail=f"Document {doc_id} not found")

        store._collection.delete(ids=ids_to_delete)
        logger.info(f"Deleted document {doc_id}: {len(ids_to_delete)} chunks removed")
        return Response(status_code=204)
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
